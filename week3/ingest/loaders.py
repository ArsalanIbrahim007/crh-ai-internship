"""Format-agnostic document loading.

Every loader yields the same Document shape, so chunking, embedding and
retrieval never learn what a PDF is. Enron is streamed straight out of the
gzipped tarball — extracting 500k small files onto NTFS costs 40 minutes and
buys nothing.
"""
from __future__ import annotations

import csv
import email
import email.utils
import hashlib
import re
import tarfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterator

from config import ENRON_TARBALL, MULTIFORMAT_DIR

# --- department inference ---------------------------------------------
# Enron maildir paths look like  maildir/lay-k/inbox/17.  The folder name and
# the address domain together give us a plausible org unit, which is what the
# RBAC layer filters on. This is heuristic and the README says so.
DEPT_RULES: list[tuple[re.Pattern, str]] = [
    (re.compile(r"legal|counsel|attorney", re.I), "Legal"),
    (re.compile(r"trading|trader|desk|gas|power", re.I), "Trading"),
    (re.compile(r"risk|raroc|credit", re.I), "Risk"),
    (re.compile(r"\bhr\b|human.?resources|benefits|staffing|recruit", re.I), "HR"),
    (re.compile(r"account|financ|tax|audit|treasur", re.I), "Finance"),
    (re.compile(r"regulat|compliance|govern|policy", re.I), "Compliance"),
    (re.compile(r"\bit\b|technolog|network|sysadmin|helpdesk", re.I), "IT"),
]
DEFAULT_DEPT = "General"

SENSITIVE = re.compile(
    r"confidential|privileged|attorney.?client|do not distribute|"
    r"internal use only|salary|compensation", re.I,
)


@dataclass
class Document:
    doc_id: str
    source: str
    fmt: str
    text: str
    title: str = ""
    author: str = ""
    department: str = DEFAULT_DEPT
    created_at: str = ""
    classification: str = "internal"     # internal | restricted
    meta: dict = field(default_factory=dict)

    @property
    def n_chars(self) -> int:
        return len(self.text)


def _doc_id(source: str, payload: str) -> str:
    h = hashlib.blake2b(f"{source}\x00{payload}".encode("utf-8", "replace"),
                        digest_size=10)
    return h.hexdigest()


def _infer_department(*signals: str) -> str:
    blob = " ".join(s for s in signals if s)
    for pattern, dept in DEPT_RULES:
        if pattern.search(blob):
            return dept
    return DEFAULT_DEPT


def _classify(text: str) -> str:
    return "restricted" if SENSITIVE.search(text[:4000]) else "internal"


def _clean(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    return text.strip()


# ----------------------------------------------------------------------
# Enron
# ----------------------------------------------------------------------
QUOTED = re.compile(
    r"^\s*(-{2,}\s*Original Message|_{5,}|>{1,}\s|On .{5,60} wrote:)", re.M
)


def _strip_quoted(body: str) -> str:
    """Drop quoted reply chains. They are duplicated across thousands of
    messages and poison both BM25 statistics and dedup."""
    m = QUOTED.search(body)
    return body[: m.start()] if m else body


def load_enron(tarball: Path = ENRON_TARBALL, limit: int | None = None,
               min_chars: int = 200) -> Iterator[Document]:
    """Stream messages out of the Enron tarball without extracting it."""
    if not tarball.exists():
        raise FileNotFoundError(
            f"{tarball} not found. Run the curl download first."
        )

    seen_hashes: set[str] = set()
    yielded = 0

    with tarfile.open(tarball, "r:gz") as tar:
        for member in tar:
            if limit is not None and yielded >= limit:
                break
            if not member.isfile():
                continue

            fh = tar.extractfile(member)
            if fh is None:
                continue
            raw = fh.read()

        # compat32 policy: headers stay raw strings. policy.default parses
            # To/From into address objects on access, and malformed real-world
            # headers crash that parser outright (CPython bpo-32178 family).
            # We store headers as metadata and never route mail, so the
            # structured form buys nothing and costs the whole run.
            try:
                msg = email.message_from_bytes(raw)
            except Exception:
                continue

            body_text = ""
            try:
                if msg.is_multipart():
                    for part in msg.walk():
                        if part.get_content_type() == "text/plain":
                            payload = part.get_payload(decode=True)
                            if payload:
                                body_text = payload.decode("utf-8", "replace")
                                break
                else:
                    payload = msg.get_payload(decode=True)
                    if payload:
                        body_text = payload.decode("utf-8", "replace")
            except Exception:
                body_text = ""

            if not body_text:
                try:
                    body_text = raw.decode("utf-8", "replace").split("\n\n", 1)[-1]
                except Exception:
                    continue

            body_text = _clean(_strip_quoted(body_text))
            if len(body_text) < min_chars:
                continue

            # near-dup guard: identical bodies are extremely common
            fp = hashlib.blake2b(body_text[:1200].encode("utf-8", "replace"),
                                 digest_size=8).hexdigest()
            if fp in seen_hashes:
                continue
            seen_hashes.add(fp)

            subject = str(msg.get("Subject") or "").strip() or "(no subject)"
            sender = str(msg.get("From") or "").strip()
            recipients = str(msg.get("To") or "").strip()
            date_hdr = str(msg.get("Date") or "").strip()

            try:
                dt = email.utils.parsedate_to_datetime(date_hdr)
                created = dt.astimezone(timezone.utc).isoformat()
            except Exception:
                created = ""

            path_parts = member.name.split("/")
            custodian = path_parts[1] if len(path_parts) > 1 else ""
            folder = path_parts[2] if len(path_parts) > 2 else ""

            dept = _infer_department(folder, subject, sender, recipients)

            yield Document(
                doc_id=_doc_id(member.name, fp),
                source=member.name,
                fmt="email",
                text=f"Subject: {subject}\n\n{body_text}",
                title=subject[:200],
                author=sender[:200],
                department=dept,
                created_at=created,
                classification=_classify(body_text),
                meta={
                    "custodian": custodian,
                    "folder": folder,
                    "recipients": recipients[:400],
                    "thread_key": re.sub(r"^\s*(re|fw|fwd)\s*:\s*", "",
                                         subject, flags=re.I).strip().lower()[:160],
                },
            )
            yielded += 1


# ----------------------------------------------------------------------
# Multi-format loaders
# ----------------------------------------------------------------------
def load_pdf(path: Path) -> Document:
    import pymupdf

    pages: list[str] = []
    with pymupdf.open(path) as doc:
        title = (doc.metadata or {}).get("title") or path.stem
        author = (doc.metadata or {}).get("author") or ""
        for i, page in enumerate(doc, start=1):
            txt = page.get_text("text").strip()
            if txt:
                # page anchors survive chunking; citations resolve to a page
                pages.append(f"[[page:{i}]]\n{txt}")
    text = _clean("\n\n".join(pages))
    return Document(
        doc_id=_doc_id(str(path), text[:1200]),
        source=str(path), fmt="pdf", text=text,
        title=str(title)[:200], author=str(author)[:200],
        department=_infer_department(path.stem, str(title)),
        created_at=datetime.fromtimestamp(path.stat().st_mtime,
                                          timezone.utc).isoformat(),
        classification=_classify(text),
        meta={"n_pages": len(pages)},
    )


def load_docx(path: Path) -> Document:
    import docx

    d = docx.Document(str(path))
    blocks = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    text = _clean("\n\n".join(blocks))
    props = d.core_properties
    title = (props.title or path.stem)
    return Document(
        doc_id=_doc_id(str(path), text[:1200]),
        source=str(path), fmt="docx", text=text,
        title=str(title)[:200], author=str(props.author or "")[:200],
        department=_infer_department(path.stem, str(title)),
        created_at=(props.created.isoformat() if props.created else ""),
        classification=_classify(text),
        meta={"n_paragraphs": len(blocks)},
    )


def load_html(path: Path) -> Document:
    from selectolax.parser import HTMLParser

    html = path.read_text(encoding="utf-8", errors="replace")
    tree = HTMLParser(html)
    for tag in tree.css("script, style, nav, footer, noscript"):
        tag.decompose()
    title_node = tree.css_first("title")
    title = title_node.text().strip() if title_node else path.stem
    body = tree.body.text(separator="\n") if tree.body else tree.text()
    text = _clean(body)
    return Document(
        doc_id=_doc_id(str(path), text[:1200]),
        source=str(path), fmt="html", text=text,
        title=title[:200],
        department=_infer_department(path.stem, title),
        created_at=datetime.fromtimestamp(path.stat().st_mtime,
                                          timezone.utc).isoformat(),
        classification=_classify(text),
    )


def load_csv(path: Path, max_rows: int = 500) -> Document:
    """Rows become sentences. A CSV chunk must read as prose or the embedding
    model has nothing to work with."""
    with path.open("r", encoding="utf-8", errors="replace", newline="") as fh:
        reader = csv.reader(fh)
        try:
            header = next(reader)
        except StopIteration:
            header = []
        lines: list[str] = []
        for i, row in enumerate(reader):
            if i >= max_rows:
                break
            pairs = [f"{h}: {v}" for h, v in zip(header, row) if v.strip()]
            if pairs:
                lines.append(f"Row {i + 1} — " + "; ".join(pairs))
    text = _clean(
        f"Table: {path.stem}\nColumns: {', '.join(header)}\n\n" + "\n".join(lines)
    )
    return Document(
        doc_id=_doc_id(str(path), text[:1200]),
        source=str(path), fmt="csv", text=text,
        title=path.stem[:200],
        department=_infer_department(path.stem, " ".join(header)),
        created_at=datetime.fromtimestamp(path.stat().st_mtime,
                                          timezone.utc).isoformat(),
        classification=_classify(text),
        meta={"columns": header[:40], "n_rows": len(lines)},
    )


LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".html": load_html,
    ".htm": load_html,
    ".csv": load_csv,
}


def load_directory(root: Path = MULTIFORMAT_DIR) -> Iterator[Document]:
    if not root.exists():
        return
    for path in sorted(root.rglob("*")):
        fn = LOADERS.get(path.suffix.lower())
        if fn is None or not path.is_file():
            continue
        try:
            doc = fn(path)
            if doc.text.strip():
                yield doc
        except Exception as exc:  # keep the pipeline alive
            from ingest import manifest
            manifest.mark_failed(_doc_id(str(path), ""), str(path),
                                 path.suffix.lstrip("."), repr(exc))