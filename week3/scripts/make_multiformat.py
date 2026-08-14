"""Render an Enron slice into PDF, DOCX, HTML and CSV.

Multi-format ingestion has to be exercised against real files with real
structure — page breaks, tables, markup — not renamed text. Emails are
grouped into thread digests first, so the generated documents are long enough
for parent-child retrieval and page-anchored citation to mean something.
"""
from __future__ import annotations

import csv
import html
import sys
from collections import defaultdict
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from config import MULTIFORMAT_DIR, MULTIFORMAT_SLICE
from ingest.loaders import load_enron

FORMATS = ("pdf", "docx", "html", "csv")


def collect_threads(target_docs: int) -> list[tuple[str, list]]:
    """Group messages by normalised subject into multi-message threads."""
    buckets: dict[str, list] = defaultdict(list)
    for doc in load_enron(limit=target_docs * 12, min_chars=300):
        key = doc.meta.get("thread_key") or doc.title.lower()
        if key and key != "(no subject)":
            buckets[key].append(doc)

    threads = [(k, v) for k, v in buckets.items() if len(v) >= 2]
    threads.sort(key=lambda kv: sum(len(d.text) for d in kv[1]), reverse=True)
    return [(k, v[:25]) for k, v in threads[:target_docs]]


def slug(text: str, n: int = 60) -> str:
    keep = "".join(c if c.isalnum() or c in " -_" else "" for c in text)
    return ("_".join(keep.split())[:n] or "thread").strip("_")


# ----------------------------------------------------------------------
def write_pdf(path: Path, key: str, msgs: list) -> None:
    import pymupdf

    doc = pymupdf.open()
    margin, width, height = 56, 595, 842
    lead = 13.5

    def new_page():
        p = doc.new_page(width=width, height=height)
        return p, margin

    page, y = new_page()
    page.insert_text((margin, y), f"Thread digest: {key[:70]}",
                     fontsize=14, fontname="hebo")
    y += 28

    for m in msgs:
        header = f"From: {m.author or 'unknown'}   Date: {m.created_at[:10] or 'n/a'}"
        blocks = [header, ""] + m.text.split("\n")
        for line in blocks:
            for seg in [line[i:i + 95] for i in range(0, max(len(line), 1), 95)]:
                if y > height - margin - lead:
                    page, y = new_page()
                page.insert_text((margin, y), seg, fontsize=9.5, fontname="helv")
                y += lead
        y += lead

    doc.set_metadata({"title": f"Thread digest: {key[:70]}",
                      "author": msgs[0].author[:100]})
    doc.save(str(path))
    doc.close()


def write_docx(path: Path, key: str, msgs: list) -> None:
    import docx

    d = docx.Document()
    d.core_properties.title = f"Thread digest: {key[:70]}"
    d.core_properties.author = msgs[0].author[:100]
    d.add_heading(f"Thread digest: {key[:70]}", level=1)

    table = d.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Sender", "Date", "Department"
    for m in msgs:
        row = table.add_row().cells
        row[0].text = m.author[:80] or "unknown"
        row[1].text = m.created_at[:10] or "n/a"
        row[2].text = m.department

    for i, m in enumerate(msgs, 1):
        d.add_heading(f"Message {i}: {m.title[:80]}", level=2)
        for para in m.text.split("\n\n"):
            if para.strip():
                d.add_paragraph(para.strip())
    d.save(str(path))


def write_html(path: Path, key: str, msgs: list) -> None:
    parts = [
        "<!doctype html><html><head><meta charset='utf-8'>",
        f"<title>Thread digest: {html.escape(key[:70])}</title>",
        "<style>body{font-family:Georgia,serif;max-width:52rem}</style>",
        "<script>console.log('nav');</script></head><body>",
        "<nav>Knowledge Platform &rsaquo; Threads</nav>",
        f"<h1>Thread digest: {html.escape(key[:70])}</h1>",
    ]
    for i, m in enumerate(msgs, 1):
        parts.append(f"<section><h2>Message {i}: {html.escape(m.title[:80])}</h2>")
        parts.append(
            f"<p><em>From {html.escape(m.author[:80] or 'unknown')} "
            f"on {html.escape(m.created_at[:10] or 'n/a')}</em></p>"
        )
        for para in m.text.split("\n\n"):
            if para.strip():
                parts.append(f"<p>{html.escape(para.strip())}</p>")
        parts.append("</section>")
    parts.append("<footer>Generated corpus artefact</footer></body></html>")
    path.write_text("\n".join(parts), encoding="utf-8")


def write_csv(path: Path, key: str, msgs: list) -> None:
    with path.open("w", newline="", encoding="utf-8") as fh:
        w = csv.writer(fh)
        w.writerow(["message_no", "sender", "recipients", "date",
                    "department", "classification", "subject", "body_excerpt"])
        for i, m in enumerate(msgs, 1):
            w.writerow([
                i, m.author[:120], m.meta.get("recipients", "")[:160],
                m.created_at[:10], m.department, m.classification,
                m.title[:120],
                m.text.replace("\n", " ")[:600],
            ])


WRITERS = {"pdf": write_pdf, "docx": write_docx,
           "html": write_html, "csv": write_csv}


def main() -> None:
    target = MULTIFORMAT_SLICE
    print(f"scanning corpus for {target} multi-message threads")
    threads = collect_threads(target)
    if not threads:
        raise SystemExit("no threads found — check the tarball")

    for fmt in FORMATS:
        (MULTIFORMAT_DIR / fmt).mkdir(parents=True, exist_ok=True)

    written = defaultdict(int)
    for i, (key, msgs) in enumerate(threads):
        fmt = FORMATS[i % len(FORMATS)]
        path = MULTIFORMAT_DIR / fmt / f"{i:04d}_{slug(key)}.{fmt}"
        try:
            WRITERS[fmt](path, key, msgs)
            written[fmt] += 1
        except Exception as exc:
            print(f"  skip {path.name}: {exc!r}")
        if (i + 1) % 100 == 0:
            print(f"  {i + 1}/{len(threads)}")

    print("\nwritten:")
    for fmt in FORMATS:
        print(f"  {fmt:5s} {written[fmt]:4d}")
    print(f"\nroot: {MULTIFORMAT_DIR}")


if __name__ == "__main__":
    main()