import csv
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import pytest

from config import ENRON_TARBALL
from ingest.loaders import (
    _classify, _infer_department, _strip_quoted, load_csv, load_docx,
    load_enron, load_html,
)


def test_department_inference():
    assert _infer_department("legal_agreements", "") == "Legal"
    assert _infer_department("", "gas trading desk") == "Trading"
    assert _infer_department("random_folder", "lunch plans") == "General"


def test_sensitive_content_is_restricted():
    assert _classify("This is attorney-client privileged material.") == "restricted"
    assert _classify("Meeting moved to 3pm.") == "internal"


def test_quoted_replies_are_stripped():
    body = "My answer is yes.\n\n-----Original Message-----\nFrom: bob"
    assert "Original Message" not in _strip_quoted(body)
    assert "My answer is yes." in _strip_quoted(body)


def test_csv_rows_become_prose(tmp_path):
    p = tmp_path / "positions.csv"
    with p.open("w", newline="", encoding="utf-8") as fh:
        w = csv.writer(fh)
        w.writerow(["desk", "volume", "region"])
        w.writerow(["gas", "1200", "west"])
    doc = load_csv(p)
    assert doc.fmt == "csv"
    assert "desk: gas" in doc.text
    assert doc.meta["n_rows"] == 1


def test_html_strips_scripts(tmp_path):
    p = tmp_path / "memo.html"
    p.write_text(
        "<html><head><title>Risk Memo</title></head>"
        "<body><script>alert(1)</script><p>Exposure is capped.</p></body></html>",
        encoding="utf-8",
    )
    doc = load_html(p)
    assert doc.title == "Risk Memo"
    assert "alert" not in doc.text
    assert "Exposure is capped." in doc.text


def test_docx_roundtrip(tmp_path):
    import docx
    d = docx.Document()
    d.add_paragraph("Compliance review findings.")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "item"
    t.rows[0].cells[1].text = "status"
    p = tmp_path / "review.docx"
    d.save(str(p))
    doc = load_docx(p)
    assert "Compliance review findings." in doc.text
    assert "item | status" in doc.text


@pytest.mark.skipif(not ENRON_TARBALL.exists(), reason="corpus not downloaded")
def test_enron_stream_yields_documents():
    docs = list(load_enron(limit=25))
    assert len(docs) == 25
    assert all(d.fmt == "email" for d in docs)
    assert all(d.text.startswith("Subject:") for d in docs)
    assert all(d.doc_id for d in docs)
    assert len({d.doc_id for d in docs}) == 25   # dedup holds